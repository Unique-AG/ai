from docx.document import Document as DocumentObject
from docxtpl import DocxTemplate
from pydantic import BaseModel


class HeadingField(BaseModel):
    text: str
    level: int = 4

    def add(self, doc: DocumentObject):
        p = doc.add_heading(self.text, level=self.level)
        return p

    def __str__(self):
        return f"HeadingField(text={self.text}, level={self.level})"


class ParagraphField(BaseModel):
    text: str
    style: str | None = None

    def add(self, doc: DocumentObject):
        p = doc.add_paragraph(self.text, style=self.style)
        return p

    def __str__(self):
        return f"ParagraphField(text={self.text}, style={self.style})"


class RunField(BaseModel):
    text: str
    italic: bool | None = False
    bold: bool | None = False

    def __str__(self):
        return f"RunField(text={self.text}, italic={self.italic}, bold={self.bold})"


class RunsField(BaseModel):
    runs: list[RunField]
    style: str | None = None

    def add(self, doc: DocumentObject):
        if not self.runs:
            return None
        p = doc.add_paragraph(style=self.style)
        for run in self.runs:
            r = p.add_run(run.text)
            if run.italic:
                r.italic = True
            if run.bold:
                r.bold = True
        return p

    def __str__(self):
        return f"RunsField(runs={self.runs}, style={self.style})"


class ContentField(BaseModel):
    contents: list[HeadingField | ParagraphField | RunsField]

    def add(self, doc: DocxTemplate):
        sd = doc.new_subdoc()
        for content in self.contents:
            # if isinstance(content, ImageField):
            #     content.download_image(self.download_path)
            #     content.add(sd) # type: ignore
            # else:
            content.add(sd)  # type: ignore
        return sd

    def __str__(self):
        return f"ContentField(contents={self.contents})"
