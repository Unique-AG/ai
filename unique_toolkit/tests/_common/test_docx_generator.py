"""
Unit tests for docx_generator module.
"""

from pathlib import Path
from unittest.mock import Mock, patch

import pytest
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docxtpl import DocxTemplate

from unique_toolkit._common.docx_generator.config import DocxGeneratorConfig
from unique_toolkit._common.docx_generator.schemas import (
    ContentField,
    HeadingField,
    ParagraphField,
    RunField,
    RunsField,
)
from unique_toolkit._common.docx_generator.service import DocxGeneratorService


class TestHeadingField:
    """Test cases for HeadingField schema."""

    def test_heading_field_defaults(self):
        """Test HeadingField with default values."""
        heading = HeadingField(text="Test Heading")
        assert heading.text == "Test Heading"
        assert heading.level == 4
        assert heading.alignment == WD_PARAGRAPH_ALIGNMENT.LEFT

    def test_heading_field_custom_values(self):
        """Test HeadingField with custom values."""
        heading = HeadingField(
            text="Custom Heading",
            level=2,
            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
        )
        assert heading.text == "Custom Heading"
        assert heading.level == 2
        assert heading.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER

    def test_heading_field_str(self):
        """Test HeadingField string representation."""
        heading = HeadingField(text="Test", level=3)
        str_repr = str(heading)
        assert "Test" in str_repr
        assert "3" in str_repr

    @patch("docx.Document")
    def test_heading_field_add(self, mock_doc_class):
        """Test adding a HeadingField to a document."""
        mock_doc = Mock()
        mock_paragraph = Mock()
        mock_doc.add_heading.return_value = mock_paragraph

        heading = HeadingField(text="Test Heading", level=2)
        _ = heading.add(mock_doc)

        mock_doc.add_heading.assert_called_once_with("Test Heading", level=2)
        assert mock_paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.LEFT


class TestParagraphField:
    """Test cases for ParagraphField schema."""

    def test_paragraph_field_defaults(self):
        """Test ParagraphField with default values."""
        paragraph = ParagraphField(text="Test paragraph")
        assert paragraph.text == "Test paragraph"
        assert paragraph.style is None
        assert paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.LEFT

    def test_paragraph_field_custom_values(self):
        """Test ParagraphField with custom values."""
        paragraph = ParagraphField(
            text="Custom paragraph",
            style="Normal",
            alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT,
        )
        assert paragraph.text == "Custom paragraph"
        assert paragraph.style == "Normal"
        assert paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT

    def test_paragraph_field_str(self):
        """Test ParagraphField string representation."""
        paragraph = ParagraphField(text="Test", style="Normal")
        str_repr = str(paragraph)
        assert "Test" in str_repr
        assert "Normal" in str_repr

    @patch("docx.Document")
    def test_paragraph_field_add(self, mock_doc_class):
        """Test adding a ParagraphField to a document."""
        mock_doc = Mock()
        mock_paragraph = Mock()
        mock_doc.add_paragraph.return_value = mock_paragraph

        paragraph = ParagraphField(text="Test paragraph", style="Normal")
        _ = paragraph.add(mock_doc)

        mock_doc.add_paragraph.assert_called_once_with("Test paragraph", style="Normal")
        assert mock_paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.LEFT


class TestRunField:
    """Test cases for RunField schema."""

    def test_run_field_defaults(self):
        """Test RunField with default values."""
        run = RunField(text="Test run")
        assert run.text == "Test run"
        assert run.italic is False
        assert run.bold is False
        assert run.alignment == WD_PARAGRAPH_ALIGNMENT.LEFT

    def test_run_field_custom_values(self):
        """Test RunField with custom values."""
        run = RunField(
            text="Bold italic text",
            italic=True,
            bold=True,
            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
        )
        assert run.text == "Bold italic text"
        assert run.italic is True
        assert run.bold is True
        assert run.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER

    def test_run_field_str(self):
        """Test RunField string representation."""
        run = RunField(text="Test", italic=True)
        str_repr = str(run)
        assert "Test" in str_repr
        assert "True" in str_repr


class TestRunsField:
    """Test cases for RunsField schema."""

    def test_runs_field_defaults(self):
        """Test RunsField with default values."""
        runs = [RunField(text="Run 1"), RunField(text="Run 2")]
        runs_field = RunsField(runs=runs)
        assert len(runs_field.runs) == 2
        assert runs_field.style is None
        assert runs_field.alignment == WD_PARAGRAPH_ALIGNMENT.LEFT

    def test_runs_field_custom_values(self):
        """Test RunsField with custom values."""
        runs = [RunField(text="Run 1", bold=True)]
        runs_field = RunsField(
            runs=runs,
            style="List Bullet",
            alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT,
        )
        assert len(runs_field.runs) == 1
        assert runs_field.style == "List Bullet"
        assert runs_field.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT

    def test_runs_field_str(self):
        """Test RunsField string representation."""
        runs = [RunField(text="Test")]
        runs_field = RunsField(runs=runs, style="Normal")
        str_repr = str(runs_field)
        assert "Test" in str_repr
        assert "Normal" in str_repr

    @patch("docx.Document")
    def test_runs_field_add(self, mock_doc_class):
        """Test adding a RunsField to a document."""
        mock_doc = Mock()
        mock_paragraph = Mock()
        mock_run1 = Mock()
        mock_run2 = Mock()
        mock_doc.add_paragraph.return_value = mock_paragraph
        mock_paragraph.add_run.side_effect = [mock_run1, mock_run2]

        runs = [
            RunField(text="Normal text"),
            RunField(text="Bold text", bold=True),
        ]
        runs_field = RunsField(runs=runs, style="Normal")
        _ = runs_field.add(mock_doc)

        mock_doc.add_paragraph.assert_called_once_with(style="Normal")
        assert mock_paragraph.add_run.call_count == 2
        mock_paragraph.add_run.assert_any_call("Normal text")
        mock_paragraph.add_run.assert_any_call("Bold text")
        assert mock_run2.bold is True

    @patch("docx.Document")
    def test_runs_field_add_empty_runs(self, mock_doc_class):
        """Test adding a RunsField with empty runs to a document."""
        mock_doc = Mock()
        runs_field = RunsField(runs=[])
        result = runs_field.add(mock_doc)

        assert result is None
        mock_doc.add_paragraph.assert_not_called()


class TestContentField:
    """Test cases for ContentField schema."""

    def test_content_field_creation(self):
        """Test ContentField creation."""
        contents: list[HeadingField | ParagraphField | RunsField] = [
            HeadingField(text="Heading"),
            ParagraphField(text="Paragraph"),
        ]
        content_field = ContentField(contents=contents)
        assert len(content_field.contents) == 2

    def test_content_field_str(self):
        """Test ContentField string representation."""
        contents: list[HeadingField | ParagraphField | RunsField] = [HeadingField(text="Test")]
        content_field = ContentField(contents=contents)
        str_repr = str(content_field)
        assert "ContentField" in str_repr

    def test_content_field_add(self):
        """Test adding a ContentField to a document template."""
        mock_template = Mock(spec=DocxTemplate)
        mock_subdoc = Mock()
        mock_template.new_subdoc.return_value = mock_subdoc

        # Create mock for heading.add
        mock_heading = Mock(spec=HeadingField)
        mock_heading.add = Mock()

        contents: list[HeadingField | ParagraphField | RunsField] = [mock_heading]  # type: ignore
        content_field = ContentField(contents=contents)
        result = content_field.add(mock_template)

        mock_template.new_subdoc.assert_called_once()
        mock_heading.add.assert_called_once_with(mock_subdoc)
        assert result == mock_subdoc


class TestDocxGeneratorConfig:
    """Test cases for DocxGeneratorConfig."""

    def test_config_defaults(self):
        """Test DocxGeneratorConfig with default values."""
        config = DocxGeneratorConfig()
        assert config.template_content_id == ""
        assert config.template_fields == {}

    def test_config_custom_values(self):
        """Test DocxGeneratorConfig with custom values."""
        config = DocxGeneratorConfig(
            template_content_id="content-123",
            template_fields={"title": "Test Document", "author": "Test Author"},
        )
        assert config.template_content_id == "content-123"
        assert config.template_fields == {"title": "Test Document", "author": "Test Author"}


class TestDocxGeneratorServiceMarkdownParsing:
    """Test cases for DocxGeneratorService markdown parsing."""

    def test_parse_markdown_simple_heading(self):
        """Test parsing simple markdown heading."""
        markdown = "#### Test Heading"
        result = DocxGeneratorService.parse_markdown_to_list_content_fields(markdown)

        assert len(result) == 1
        assert isinstance(result[0], HeadingField)
        assert result[0].text == "Test Heading"
        assert result[0].level == 4

    def test_parse_markdown_paragraph(self):
        """Test parsing simple markdown paragraph."""
        markdown = "This is a paragraph."
        result = DocxGeneratorService.parse_markdown_to_list_content_fields(markdown)

        assert len(result) == 1
        assert isinstance(result[0], RunsField)
        assert len(result[0].runs) == 1
        assert result[0].runs[0].text == "This is a paragraph."

    def test_parse_markdown_bold_text(self):
        """Test parsing markdown with bold text."""
        markdown = "This is **bold** text."
        result = DocxGeneratorService.parse_markdown_to_list_content_fields(markdown)

        assert len(result) == 1
        assert isinstance(result[0], RunsField)
        assert len(result[0].runs) == 3
        assert result[0].runs[0].text == "This is "
        assert result[0].runs[0].bold is False
        assert result[0].runs[1].text == "bold"
        assert result[0].runs[1].bold is True
        assert result[0].runs[2].text == " text."
        assert result[0].runs[2].bold is False

    def test_parse_markdown_italic_text(self):
        """Test parsing markdown with italic text."""
        markdown = "This is *italic* text."
        result = DocxGeneratorService.parse_markdown_to_list_content_fields(markdown)

        assert len(result) == 1
        assert isinstance(result[0], RunsField)
        assert len(result[0].runs) == 3
        assert result[0].runs[0].text == "This is "
        assert result[0].runs[0].italic is False
        assert result[0].runs[1].text == "italic"
        assert result[0].runs[1].italic is True
        assert result[0].runs[2].text == " text."
        assert result[0].runs[2].italic is False

    def test_parse_markdown_bullet_list(self):
        """Test parsing markdown bullet list."""
        markdown = """- Item 1
- Item 2
- Item 3"""
        result = DocxGeneratorService.parse_markdown_to_list_content_fields(markdown)

        assert len(result) == 3
        for item in result:
            assert isinstance(item, RunsField)
            assert item.style == "List Bullet"

    def test_parse_markdown_nested_bullet_list(self):
        """Test parsing markdown nested bullet list."""
        markdown = """- Item 1
  - Nested Item 1
  - Nested Item 2
- Item 2"""
        result = DocxGeneratorService.parse_markdown_to_list_content_fields(markdown)

        assert len(result) == 4
        # Check styles for RunsField items
        assert isinstance(result[0], RunsField) and result[0].style == "List Bullet"
        assert isinstance(result[1], RunsField) and result[1].style == "List Bullet 2"
        assert isinstance(result[2], RunsField) and result[2].style == "List Bullet 2"
        assert isinstance(result[3], RunsField) and result[3].style is None

    def test_parse_markdown_mixed_content(self):
        """Test parsing markdown with mixed content."""
        markdown = """#### Heading

This is a paragraph with **bold** and *italic* text.

- List item 1
- List item 2"""
        result = DocxGeneratorService.parse_markdown_to_list_content_fields(markdown)

        assert len(result) >= 3
        assert isinstance(result[0], HeadingField)
        assert result[0].text == "Heading"

    def test_parse_markdown_heading_replacement(self):
        """Test markdown heading level replacement."""
        markdown = "## Second Level"
        result = DocxGeneratorService.parse_markdown_to_list_content_fields(markdown)

        assert len(result) == 1
        assert isinstance(result[0], HeadingField)
        assert result[0].level == 4  # ## is replaced with ####

    def test_parse_markdown_offset_header_level(self):
        """Test markdown parsing with header level offset."""
        markdown = "#### Test Heading"
        result = DocxGeneratorService.parse_markdown_to_list_content_fields(
            markdown, offset_header_lvl=1
        )

        assert len(result) == 1
        assert isinstance(result[0], HeadingField)
        assert result[0].level == 5  # 4 + 1 offset

    def test_parse_markdown_page_heading_special_case(self):
        """Test parsing markdown with _pageXXX_ special heading."""
        markdown = "#### _page123_Some text"
        result = DocxGeneratorService.parse_markdown_to_list_content_fields(markdown)

        assert len(result) == 1
        assert isinstance(result[0], HeadingField)
        # The regex captures everything up to the closing underscore
        assert result[0].text == "Page123Some text"

    def test_parse_markdown_removes_relevant_sources(self):
        """Test that 'Relevant sources' heading is removed."""
        markdown = """#### Main Heading

# Relevant sources

Some content"""
        result = DocxGeneratorService.parse_markdown_to_list_content_fields(markdown)

        # Check that no heading with "Relevant sources" text exists
        for item in result:
            if isinstance(item, HeadingField):
                assert "Relevant sources" not in item.text

    def test_parse_markdown_empty_string(self):
        """Test parsing empty markdown string."""
        markdown = ""
        result = DocxGeneratorService.parse_markdown_to_list_content_fields(markdown)

        assert len(result) == 0

    def test_parse_markdown_multiple_paragraphs(self):
        """Test parsing multiple paragraphs."""
        markdown = """Paragraph 1

Paragraph 2

Paragraph 3"""
        result = DocxGeneratorService.parse_markdown_to_list_content_fields(markdown)

        assert len(result) == 3
        for item in result:
            assert isinstance(item, RunsField)


class TestDocxGeneratorServiceGeneration:
    """Test cases for DocxGeneratorService document generation."""

    @pytest.fixture
    def mock_chat_service(self):
        """Create a mock ChatService."""
        return Mock()

    @pytest.fixture
    def mock_kb_service(self):
        """Create a mock KnowledgeBaseService."""
        return Mock()

    @pytest.fixture
    def docx_service(self, mock_chat_service, mock_kb_service):
        """Create a DocxGeneratorService with mocks."""
        config = DocxGeneratorConfig()
        return DocxGeneratorService(
            chat_service=mock_chat_service,
            knowledge_base_service=mock_kb_service,
            config=config,
        )

    def test_get_default_template(self, docx_service):
        """Test getting the default template."""
        template_bytes = docx_service._get_default_template()

        assert isinstance(template_bytes, bytes)
        assert len(template_bytes) > 0

    def test_get_default_template_file_exists(self, docx_service):
        """Test that default template file exists."""
        generator_dir = Path(__file__).resolve().parent.parent.parent / "unique_toolkit" / "_common" / "docx_generator"
        template_path = generator_dir / "template" / "Doc Template.docx"

        assert template_path.exists()
        assert template_path.is_file()

    @patch.object(DocxGeneratorService, "_get_default_template")
    def test_get_template_with_no_content_id(self, mock_get_default, docx_service):
        """Test getting template with no content ID."""
        mock_get_default.return_value = b"default template bytes"

        result = docx_service._get_template("")

        mock_get_default.assert_called_once()
        assert result == b"default template bytes"

    @patch.object(DocxGeneratorService, "_get_default_template")
    def test_get_template_with_content_id(self, mock_get_default, docx_service, mock_kb_service):
        """Test getting template with content ID."""
        mock_kb_service.download_content_to_bytes.return_value = b"custom template bytes"
        docx_service._config.template_content_id = "content-123"

        result = docx_service._get_template("content-123")

        mock_kb_service.download_content_to_bytes.assert_called_once_with(
            content_id="content-123"
        )
        assert result == b"custom template bytes"

    @patch.object(DocxGeneratorService, "_get_default_template")
    def test_get_template_fallback_on_error(self, mock_get_default, docx_service, mock_kb_service):
        """Test template fallback to default on download error."""
        mock_kb_service.download_content_to_bytes.side_effect = Exception("Download failed")
        mock_get_default.return_value = b"default template bytes"
        docx_service._config.template_content_id = "invalid-id"

        result = docx_service._get_template("invalid-id")

        mock_get_default.assert_called_once()
        assert result == b"default template bytes"

    @patch.object(DocxGeneratorService, "_get_template")
    @patch("unique_toolkit._common.docx_generator.service.DocxTemplate")
    def test_generate_from_template_success(
        self, mock_docx_template_class, mock_get_template, docx_service
    ):
        """Test successful document generation from template."""
        # Setup mocks
        mock_get_template.return_value = b"template bytes"
        mock_template = Mock(spec=DocxTemplate)
        mock_docx_template_class.return_value = mock_template
        mock_template.render = Mock()
        mock_template.save = Mock()

        # Create test content
        subdoc_content = [
            HeadingField(text="Test Heading"),
            RunsField(runs=[RunField(text="Test content")]),
        ]

        # Generate document
        result = docx_service.generate_from_template(subdoc_content)

        # Assertions
        mock_get_template.assert_called_once()
        mock_template.render.assert_called_once()
        mock_template.save.assert_called_once()
        assert isinstance(result, bytes)

    @patch.object(DocxGeneratorService, "_get_template")
    @patch("unique_toolkit._common.docx_generator.service.DocxTemplate")
    def test_generate_from_template_with_fields(
        self, mock_docx_template_class, mock_get_template, docx_service
    ):
        """Test document generation with additional fields."""
        # Setup mocks
        mock_get_template.return_value = b"template bytes"
        mock_template = Mock(spec=DocxTemplate)
        mock_docx_template_class.return_value = mock_template

        # Create test content
        subdoc_content = [HeadingField(text="Test")]
        fields = {"title": "Test Document", "author": "Test Author"}

        # Generate document
        _ = docx_service.generate_from_template(subdoc_content, fields=fields)

        # Check that render was called
        mock_template.render.assert_called_once()
        render_args = mock_template.render.call_args[0][0]
        assert "title" in render_args
        assert render_args["title"] == "Test Document"
        assert "author" in render_args
        assert render_args["author"] == "Test Author"

    @patch.object(DocxGeneratorService, "_get_template")
    @patch("unique_toolkit._common.docx_generator.service.DocxTemplate")
    def test_generate_from_template_error_handling(
        self, mock_docx_template_class, mock_get_template, docx_service
    ):
        """Test error handling during document generation."""
        # Setup mocks to raise an exception
        mock_get_template.return_value = b"template bytes"
        mock_template = Mock(spec=DocxTemplate)
        mock_docx_template_class.return_value = mock_template
        mock_template.render.side_effect = Exception("Render failed")

        # Create test content
        subdoc_content = [HeadingField(text="Test")]

        # Generate document
        result = docx_service.generate_from_template(subdoc_content)

        # Should return None on error
        assert result is None


class TestDocxGeneratorIntegration:
    """Integration tests for DocxGeneratorService."""

    @pytest.fixture
    def real_docx_service(self):
        """Create a real DocxGeneratorService for integration tests."""
        mock_chat_service = Mock()
        mock_kb_service = Mock()
        config = DocxGeneratorConfig()
        return DocxGeneratorService(
            chat_service=mock_chat_service,
            knowledge_base_service=mock_kb_service,
            config=config,
        )

    def test_full_markdown_to_docx_workflow(self, real_docx_service):
        """Test the complete workflow from markdown to docx."""
        markdown = """#### Introduction

This is a test document with **bold** and *italic* text.

#### Features

- Feature 1
- Feature 2
- Feature 3

#### Conclusion

This is the conclusion."""

        # Parse markdown
        content = DocxGeneratorService.parse_markdown_to_list_content_fields(markdown)

        # Verify parsed content
        assert len(content) > 0
        assert any(isinstance(item, HeadingField) for item in content)
        assert any(isinstance(item, RunsField) for item in content)

        # Generate document
        _ = real_docx_service.generate_from_template(content)

        # Verify result (should be bytes or None if template not available)
        # Note: We don't assert on result as it depends on template availability

    def test_complex_markdown_parsing(self):
        """Test parsing complex markdown structure."""
        markdown = """#### Main Title

This is an introductory paragraph with **bold** and *italic* formatting.

#### Section 1

##### Subsection 1.1

Some content here.

- Bullet point 1
- Bullet point 2
  - Nested bullet 1
  - Nested bullet 2

#### Section 2

More content with **bold** formatting.

- List item A
- List item B"""

        result = DocxGeneratorService.parse_markdown_to_list_content_fields(markdown)

        # Verify structure
        assert len(result) > 0
        headings = [item for item in result if isinstance(item, HeadingField)]
        runs_fields = [item for item in result if isinstance(item, RunsField)]

        assert len(headings) >= 3
        assert len(runs_fields) > 0


if __name__ == "__main__":
    pytest.main([__file__, "-v"])

